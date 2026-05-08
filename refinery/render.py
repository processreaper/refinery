"""Render markdown text to PDF or DOCX (built-in fallbacks)."""

from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import escape as _xml_escape

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _sanitize(text: str) -> str:
    return _CONTROL_RE.sub("", text)


def _markdown_to_reportlab_html(text: str) -> str:
    safe = _xml_escape(_sanitize(text))
    safe = _BOLD_RE.sub(r"<b>\1</b>", safe)
    safe = _ITALIC_RE.sub(r"<i>\1</i>", safe)
    return safe


def render_markdown_to_pdf(markdown_text: str, dest: Path) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    flowables: list = []
    for line in _sanitize(markdown_text).split("\n"):
        if not line.strip():
            flowables.append(Spacer(1, 0.12 * inch))
            continue

        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            style = styles[f"Heading{min(level, 6)}"]
            flowables.append(Paragraph(_markdown_to_reportlab_html(m.group(2)), style))
            continue

        m = _BULLET_RE.match(line)
        if m:
            flowables.append(
                Paragraph("• " + _markdown_to_reportlab_html(m.group(1)), styles["BodyText"])
            )
            continue

        flowables.append(
            Paragraph(_markdown_to_reportlab_html(line), styles["BodyText"])
        )

    doc = SimpleDocTemplate(
        str(dest),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Refinery document",
    )
    doc.build(flowables)


def _add_emphasis_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_markdown_to_docx(markdown_text: str, dest: Path) -> None:
    from docx import Document

    doc = Document()
    for line in _sanitize(markdown_text).split("\n"):
        if not line.strip():
            continue

        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 9))
            continue

        m = _BULLET_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_emphasis_runs(p, m.group(1))
            continue

        p = doc.add_paragraph()
        _add_emphasis_runs(p, line)

    doc.save(str(dest))
