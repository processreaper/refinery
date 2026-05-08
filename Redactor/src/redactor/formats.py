"""Format adapters: read text from a document, write redacted text back."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

Redact = Callable[[str], str]

# Output formats `redact_file` can produce when given output_format=<...>.
# `original` (handled separately) preserves the input's structure.
# Always-available: txt, md, pdf, docx (built-in renderers).
# Pandoc-only: html, rtf, odt, epub.
OUTPUT_FORMATS = {"txt", "md", "pdf", "docx", "html", "rtf", "odt", "epub"}
PANDOC_ONLY_FORMATS = {"html", "rtf", "odt", "epub"}


def _render_output(redacted: str, dest: Path, output_format: str) -> None:
    """Write `redacted` (markdown text) to `dest` in the requested format.

    Prefers pandoc when available — it produces nicer output, especially
    for DOCX. Falls back to the built-in renderers (reportlab / python-docx)
    for pdf and docx. Pandoc-only formats raise if pandoc isn't installed.
    """
    if output_format == "txt":
        dest.write_text(redacted, encoding="utf-8")
        return
    if output_format == "md":
        dest.write_text(redacted, encoding="utf-8")
        return

    from redactor import pandoc_render

    if pandoc_render.can_render(output_format):
        pandoc_render.render(redacted, dest, output_format)
        return

    if output_format in PANDOC_ONLY_FORMATS:
        if not pandoc_render.is_pandoc_available():
            raise RuntimeError(
                f"--output-format {output_format} requires pandoc. "
                "Install it (e.g. `brew install pandoc` on macOS) and retry."
            )
        raise RuntimeError(
            f"pandoc cannot produce {output_format} on this machine"
        )

    # Fall back to built-in renderers for pdf / docx.
    if output_format == "pdf":
        from redactor.render import render_markdown_to_pdf

        render_markdown_to_pdf(redacted, dest)
        return
    if output_format == "docx":
        from redactor.render import render_markdown_to_docx

        render_markdown_to_docx(redacted, dest)
        return

    raise RuntimeError(f"unknown output format: {output_format}")


def detect_format(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".eml":
        return "eml"
    if suffix == ".msg":
        return "msg"
    return "text"


def redact_text_file(src: Path, dest: Path, redact: Redact) -> None:
    text = src.read_text(encoding="utf-8")
    dest.write_text(redact(text), encoding="utf-8")


def extract_pdf_text(src: Path) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(src) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def redact_pdf_file(src: Path, dest: Path, redact: Redact) -> None:
    """Extract text from each page, redact, and render a fresh PDF.

    The original PDF's layout (images, fonts, columns, tables) is not
    carried through &mdash; the result is a newly typeset PDF containing
    the redacted text. Use --output-format txt|md to get plain text out.
    """
    text = extract_pdf_text(src)
    _render_output(redact(text), dest, "pdf")


def redact_docx_file(src: Path, dest: Path, redact: Redact) -> None:
    """Redact a .docx in place by paragraph.

    Run-level formatting within a paragraph is collapsed onto the first run,
    because entity boundaries don't align with run boundaries. Paragraph-level
    structure (lists, tables, headings) is preserved.
    """
    from docx import Document

    doc = Document(str(src))

    def _redact_paragraph(paragraph) -> None:
        if not paragraph.text:
            return
        new_text = redact(paragraph.text)
        if new_text == paragraph.text:
            return
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    for paragraph in doc.paragraphs:
        _redact_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _redact_paragraph(paragraph)

    doc.save(str(dest))


def redact_file(
    src: Path,
    dest: Path,
    redact: Redact,
    fmt: str | None = None,
    output_format: str = "original",
) -> None:
    """Redact `src` into `dest`.

    `output_format` controls the rendering of the redacted output:
    - "original" (default): keep the input format (DOCX→DOCX, EML→EML, etc.)
    - "txt": flatten to plain UTF-8 text, regardless of input format
    - "md":  flatten to markdown (DOCX styles → headings/lists; emails get
             a header block followed by the body)
    """
    fmt = fmt or detect_format(src)

    if output_format in OUTPUT_FORMATS:
        from redactor.extract import extract_for_output

        # Rendered formats (anything that's not txt) keep heading/list/email
        # structure, so we extract as markdown. txt is the only flat case.
        as_markdown = output_format != "txt"
        text = extract_for_output(src, fmt, markdown=as_markdown)
        redacted = redact(text)
        _render_output(redacted, dest, output_format)
        return

    if fmt == "pdf":
        redact_pdf_file(src, dest, redact)
    elif fmt == "docx":
        redact_docx_file(src, dest, redact)
    elif fmt == "eml":
        from redactor.email_format import redact_eml_file
        redact_eml_file(src, dest, redact)
    elif fmt == "msg":
        from redactor.email_format import redact_msg_file
        redact_msg_file(src, dest, redact)
    else:
        redact_text_file(src, dest, redact)
