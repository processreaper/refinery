"""Per-format text/markdown extractors used when --output-format is txt or md."""

from __future__ import annotations

import html as _html
import re
from email import message_from_bytes
from email.message import EmailMessage
from email.policy import default as default_policy
from pathlib import Path

EMAIL_HEADER_ORDER = ["From", "To", "Cc", "Bcc", "Reply-To", "Subject", "Date"]


def _strip_html(html: str) -> str:
    text = re.sub(r"<\s*br\s*/?\s*>", "\n", html, flags=re.I)
    text = re.sub(r"</\s*p\s*>", "\n\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return _html.unescape(text).strip()


def extract_text_passthrough(src: Path) -> str:
    return src.read_text(encoding="utf-8")


def extract_pdf(src: Path) -> str:
    from redactor.formats import extract_pdf_text

    return extract_pdf_text(src)


def extract_docx_text(src: Path) -> str:
    from docx import Document

    doc = Document(str(src))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return "\n\n".join(parts)


def extract_docx_markdown(src: Path) -> str:
    from docx import Document

    doc = Document(str(src))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text
        if not text.strip():
            continue
        style = (p.style.name if p.style else "").lower()
        if style == "title":
            parts.append(f"# {text}")
        elif style.startswith("heading 1"):
            parts.append(f"# {text}")
        elif style.startswith("heading 2"):
            parts.append(f"## {text}")
        elif style.startswith("heading 3"):
            parts.append(f"### {text}")
        elif style.startswith("heading 4"):
            parts.append(f"#### {text}")
        elif style.startswith("heading 5"):
            parts.append(f"##### {text}")
        elif style.startswith("heading 6"):
            parts.append(f"###### {text}")
        elif "bullet" in style or "list" in style:
            parts.append(f"- {text}")
        else:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return "\n\n".join(parts)


def _email_body(msg: EmailMessage) -> str:
    if msg.is_multipart():
        text_plain: str | None = None
        text_html: str | None = None
        for part in msg.walk():
            if part.is_multipart():
                continue
            ctype = part.get_content_type()
            try:
                content = part.get_content()
            except (LookupError, KeyError, ValueError):
                continue
            if not isinstance(content, str):
                continue
            if ctype == "text/plain" and text_plain is None:
                text_plain = content
            elif ctype == "text/html" and text_html is None:
                text_html = content
        if text_plain:
            return text_plain.strip()
        if text_html:
            return _strip_html(text_html)
        return ""

    ctype = msg.get_content_type()
    try:
        content = msg.get_content()
    except (LookupError, KeyError, ValueError):
        return ""
    if not isinstance(content, str):
        return ""
    if ctype == "text/html":
        return _strip_html(content)
    return content.strip()


def _email_to_string(msg: EmailMessage, markdown: bool) -> str:
    lines: list[str] = []
    for h in EMAIL_HEADER_ORDER:
        v = msg[h]
        if v:
            lines.append(f"**{h}:** {v}" if markdown else f"{h}: {v}")
    if not lines:
        return _email_body(msg)
    if markdown:
        lines.append("")
        lines.append("---")
        lines.append("")
    else:
        lines.append("")
    lines.append(_email_body(msg))
    return "\n".join(lines)


def extract_eml(src: Path, markdown: bool = False) -> str:
    msg = message_from_bytes(src.read_bytes(), policy=default_policy)
    return _email_to_string(msg, markdown)


def extract_msg(src: Path, markdown: bool = False) -> str:
    from redactor.email_format import _msg_to_email_message

    msg = _msg_to_email_message(src)
    return _email_to_string(msg, markdown)


def extract_for_output(src: Path, fmt: str, markdown: bool) -> str:
    """Return a plain-text or markdown rendering of `src` based on its format."""
    if fmt == "pdf":
        return extract_pdf(src)
    if fmt == "docx":
        return extract_docx_markdown(src) if markdown else extract_docx_text(src)
    if fmt == "eml":
        return extract_eml(src, markdown=markdown)
    if fmt == "msg":
        return extract_msg(src, markdown=markdown)
    return extract_text_passthrough(src)
